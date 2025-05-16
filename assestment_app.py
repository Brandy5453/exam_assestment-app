'''
import streamlit as st
import os
import pandas as pd
from docx import Document
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_absolute_error

# ----------------------- STEP 1: EXTRACT FEATURES FROM DOCX ------------------------

def extract_docx_features(doc_path):
    """Extracts text, formatting, images, and text boxes from a DOCX file."""
    doc = Document(doc_path)
    extracted_data = {
        "text": [],
        "bold": [],
        "font_size": [],
        "images": 0,
        "text_boxes": 0
    }

    # Extract paragraphs and formatting
    for para in doc.paragraphs:
        extracted_data["text"].append(para.text)
        
        # Formatting analysis
        bold_count = sum(1 for run in para.runs if run.bold)
        font_sizes = [run.font.size.pt if run.font.size else 12 for run in para.runs]
        
        extracted_data["bold"].append(bold_count)
        extracted_data["font_size"].append(sum(font_sizes) / len(font_sizes) if font_sizes else 12)

    # Count images and text boxes
    for shape in doc.inline_shapes:
        extracted_data["images"] += 1

    return extracted_data


# ----------------------- STEP 2: COMPARE WITH EXPECTED ANSWERS ------------------------

EXPECTED_ANSWERS = {
    "text": [
        "Exciting Internship Opportunity for Aspiring Social Media Influencers!",
        "Strong understanding of various social media platforms (Instagram, TikTok, YouTube, etc.)",
        "Creativity and a knack for storytelling through captivating content",
        "Excellent communication skills, both written and verbal",
        "Ability to collaborate effectively with team members and influencers"
    ],
    "bold": [0, 1, 1, 1, 1],  
    "font_size": [44, 11, 11, 11, 11],
    "images": 1,
    "text_boxes": 1
}

def calculate_score(extracted_data):
    """Compares extracted features with expected values and assigns marks."""
    total_score = 0

    # Check text presence
    for expected_text in EXPECTED_ANSWERS["text"]:
        if any(expected_text in text for text in extracted_data["text"]):
            total_score += 5

    # Check formatting
    total_score += sum(
        4 for i in range(len(EXPECTED_ANSWERS["bold"])) 
        if extracted_data["bold"][i] == EXPECTED_ANSWERS["bold"][i]
    )
    
    total_score += sum(
        4 for i in range(len(EXPECTED_ANSWERS["font_size"])) 
        if round(extracted_data["font_size"][i]) == EXPECTED_ANSWERS["font_size"][i]
    )

    # Check images
    if extracted_data["images"] >= EXPECTED_ANSWERS["images"]:
        total_score += 10

    return total_score


# ----------------------- STEP 3: STREAMLIT APP ------------------------

st.title("📄 Word Document Grader")
st.write("Upload a `.docx` file to get an automated assessment score.")

uploaded_file = st.file_uploader("Upload your Word Document", type=["docx"])

if uploaded_file:
    # Save uploaded file
    file_path = f"temp_{uploaded_file.name}"
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Extract features
    extracted_data = extract_docx_features(file_path)

    # Calculate score
    score = calculate_score(extracted_data)

    # Display results
    st.subheader("📊 Extracted Features:")
    st.json(extracted_data)

    st.subheader(f"✅ Predicted Score: {score}/80")

    # Delete temp file
    os.remove(file_path)
'''

import streamlit as st
import pandas as pd
from docx import Document
import spacy
from transformers import pipeline
import re
from openpyxl import load_workbook
from difflib import SequenceMatcher
import os

# Set page configuration as the first Streamlit command
st.set_page_config(page_title="University of South Africa - EUP Tool Grading System", layout="wide")

# Initialize NLP models
try:
    nlp = spacy.load("en_core_web_sm")
except:
    st.warning("Couldn't load spaCy model, using simpler text processing")
    nlp = None

try:
    zero_shot_classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
except:
    st.warning("Couldn't load zero-shot classifier, using simpler matching")
    zero_shot_classifier = None

# Function to load rubric from Excel
def load_rubric(file):
    try:
        wb = load_workbook(filename=file, read_only=True)
        sheet = wb.active
        
        rubric_data = []
        for row in sheet.iter_rows(values_only=True):
            if row and row[0] and isinstance(row[0], str) and not row[0].startswith("Criteria"):
                rubric_data.append({
                    'Criteria': row[0],
                    'Total Points': float(row[1]) if isinstance(row[1], (int, float)) else 0.0,
                    'Excellent': row[2] if len(row) > 2 else "",
                    'Good': row[3] if len(row) > 3 else "",
                    'Poor': row[4] if len(row) > 4 else ""
                })
        return pd.DataFrame(rubric_data)
    except Exception as e:
        st.error(f"Error loading rubric: {str(e)}")
        return None

# Function to extract features from a Word document
def extract_document_features(doc_file):
    try:
        doc = Document(doc_file)
        features = {
            'text': [para.text for para in doc.paragraphs],
            'raw_text': "\n".join(para.text for para in doc.paragraphs),
            'has_images': len(doc.inline_shapes) > 0,
            'images_count': len(doc.inline_shapes),
            'headings': [para.text for para in doc.paragraphs if para.style.name.startswith('Heading')],
            'lists': [],
            'formatting': {'bold': [], 'italic': [], 'underlined': [], 'colored_text': [], 'font_sizes': []},
            'tables': len(doc.tables),
            'sections': len(doc.sections),
            'bold_count': 0,
            'text_boxes': 0
        }

        # Detect lists
        for para in doc.paragraphs:
            if para.style.name.lower().startswith(('list', 'heading')) or para.text.strip().startswith(('•', '-', '*', '1.', 'a)', 'I.')):
                features['lists'].append(para.text)

        # Detect formatting and other features
        for para in doc.paragraphs:
            bold_count = sum(1 for run in para.runs if run.bold)
            font_sizes = [run.font.size.pt if run.font.size else 12 for run in para.runs]
            avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12

            features['bold_count'] += bold_count
            features['formatting']['font_sizes'].append(avg_font_size)

            for run in para.runs:
                if run.bold:
                    features['formatting']['bold'].append(run.text)
                if run.italic:
                    features['formatting']['italic'].append(run.text)
                if run.underline:
                    features['formatting']['underlined'].append(run.text)
                if hasattr(run.font, 'color') and run.font.color.rgb is not None:
                    features['formatting']['colored_text'].append(run.text)

        return features
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return None

# Predefined expected answers for simple grading (from the first script)
EXPECTED_ANSWERS = {
    "text": [
        "Exciting Internship Opportunity for Aspiring Social Media Influencers!",
        "Strong understanding of various social media platforms (Instagram, TikTok, YouTube, etc.)",
        "Creativity and a knack for storytelling through captivating content",
        "Excellent communication skills, both written and verbal",
        "Ability to collaborate effectively with team members and influencers"
    ],
    "bold": [0, 1, 1, 1, 1],
    "font_size": [44, 11, 11, 11, 11],
    "images": 1,
    "text_boxes": 1
}

# Simple grading function (adapted from the first script)
def calculate_simple_score(features):
    total_score = 0

    # Check text presence
    for expected_text in EXPECTED_ANSWERS["text"]:
        if any(expected_text in text for text in features['text']):
            total_score += 5

    # Check formatting (bold and font size)
    bold_matches = sum(1 for i in range(min(len(features['text'], len(EXPECTED_ANSWERS['bold'])))) 
                       if features['bold_count'] >= EXPECTED_ANSWERS['bold'][i])
    font_size_matches = sum(1 for i in range(min(len(features['formatting']['font_sizes'], len(EXPECTED_ANSWERS['font_size'])))) 
                            if round(features['formatting']['font_sizes'][i]) == EXPECTED_ANSWERS['font_size'][i])
    
    total_score += bold_matches * 4
    total_score += font_size_matches * 4

    # Check images
    if features['images_count'] >= EXPECTED_ANSWERS['images']:
        total_score += 10

    return total_score

# Similarity score for rubric-based grading
def similarity_score(a, b):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

# Match criteria using NLP or similarity
def match_criteria(student_features, criteria_text):
    if not student_features or not criteria_text:
        return 0.0
    
    # Exact match
    if criteria_text.lower() in student_features['raw_text'].lower():
        return 1.0
    
    # Match with headings
    for heading in student_features['headings']:
        if similarity_score(heading, criteria_text) > 0.7:
            return 1.0
    
    # Use zero-shot classification if available
    if zero_shot_classifier:
        try:
            result = zero_shot_classifier(student_features['raw_text'], [criteria_text])
            return result['scores'][0]
        except:
            pass
    
    # Fallback to text similarity
    return similarity_score(student_features['raw_text'], criteria_text)

# Grade based on rubric
def grade_assignment(student_features, rubric_df):
    if student_features is None:
        return [], 0.0, 0.0
    
    results = []
    total_possible = rubric_df['Total Points'].sum()
    total_earned = 0.0
    
    for _, row in rubric_df.iterrows():
        criteria = row['Criteria']
        max_points = float(row['Total Points'])
        
        if max_points <= 0:
            results.append({
                'Criteria': criteria,
                'Points Possible': 0,
                'Points Earned': 0,
                'Feedback': "No points allocated",
                'Match %': "N/A"
            })
            continue
        
        if "insert" in criteria.lower() or "add" in criteria.lower():
            match_score = match_criteria(student_features, criteria)
            points = max_points * match_score
            feedback = f"Content match: {match_score:.1%}"
        
        elif "format" in criteria.lower() or "style" in criteria.lower():
            points = 0
            if "bold" in criteria.lower() and student_features['formatting']['bold']:
                points = max_points * 0.8
            elif "italic" in criteria.lower() and student_features['formatting']['italic']:
                points = max_points * 0.8
            elif "color" in criteria.lower() and student_features['formatting']['colored_text']:
                points = max_points * 0.5
            elif "underline" in criteria.lower() and student_features['formatting']['underlined']:
                points = max_points * 0.5
            else:
                points = max_points * 0.3
            feedback = "Formatting assessed"
        
        elif "image" in criteria.lower() or "picture" in criteria.lower():
            points = max_points if student_features['has_images'] else 0
            feedback = "Image requirements assessed"
        
        elif "table" in criteria.lower():
            points = max_points if student_features['tables'] > 0 else 0
            feedback = "Table requirements assessed"
        
        else:
            match_score = match_criteria(student_features, criteria)
            points = max_points * match_score
            feedback = f"Generic criteria match: {match_score:.1%}"
        
        results.append({
            'Criteria': criteria,
            'Points Possible': max_points,
            'Points Earned': round(points, 1),
            'Feedback': feedback,
            'Match %': f"{min(100, int(points/max_points*100))}%" if max_points > 0 else "N/A"
        })
        total_earned += points
    
    return results, total_earned, total_possible

# Display grading results
def display_results(results, total_earned, total_possible, student_info, simple_score):
    st.subheader("Grading Results")
    st.write(f"**Student:** {student_info['name']} | **ID:** {student_info['id']}")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Rubric Score", f"{total_earned:.1f}/{total_possible:.1f}")
    with col2:
        percentage = (total_earned/total_possible)*100 if total_possible > 0 else 0
        st.metric("Rubric Percentage", f"{percentage:.1f}%")
    with col3:
        st.metric("Simple Score", f"{simple_score}/80")
    
    st.subheader("Detailed Rubric Assessment")
    results_df = pd.DataFrame(results)
    st.dataframe(results_df.style
                 .highlight_max(axis=0, subset=['Points Earned'])
                 .format({'Points Possible': '{:.1f}', 'Points Earned': '{:.1f}'}))

    # Visualizations
    st.subheader("Performance Breakdown")
    if not results_df.empty:
        viz_df = results_df.copy()
        viz_df['Match Numeric'] = viz_df['Match %'].apply(
            lambda x: float(x.replace('%','')) if x != 'N/A' else 0)
        
        col1, col2 = st.columns(2)
        with col1:
            st.bar_chart(viz_df.set_index('Criteria')['Points Earned'])
        with col2:
            st.bar_chart(viz_df.set_index('Criteria')['Match Numeric'])

# Main function
def main():
    st.title("University of South Africa - EUP Tool Grading System")

    # Initialize session state
    if 'graded_assignments' not in st.session_state:
        st.session_state.graded_assignments = []

    # File upload section
    with st.expander("Upload Files", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            rubric_file = st.file_uploader("Upload Rubric (Excel)", type=["xlsx"])
        with col2:
            assignment_file = st.file_uploader("Upload Student Assignment (Word)", type=["docx"])

    # Student info section
    student_info = {}
    if rubric_file and assignment_file:
        with st.expander("Student Information", expanded=True):
            student_info['name'] = st.text_input("Student Name", key="student_name")
            student_info['id'] = st.text_input("Student ID", key="student_id")

    # Process grading
    if rubric_file and assignment_file and student_info.get('name') and student_info.get('id'):
        # Save uploaded file temporarily
        file_path = f"temp_{assignment_file.name}"
        with open(file_path, "wb") as f:
            f.write(assignment_file.getbuffer())

        with st.spinner("Analyzing assignment..."):
            # Extract features
            student_features = extract_document_features(file_path)

            if student_features is not None:
                # Simple grading (from first script)
                simple_score = calculate_simple_score(student_features)

                # Rubric-based grading
                rubric_df = load_rubric(rubric_file)
                if rubric_df is not None:
                    results, total_earned, total_possible = grade_assignment(student_features, rubric_df)

                    # Display results
                    display_results(results, total_earned, total_possible, student_info, simple_score)

                    # Add to session state
                    st.session_state.graded_assignments.append({
                        'student_info': student_info,
                        'results': results,
                        'total_earned': total_earned,
                        'total_possible': total_possible,
                        'simple_score': simple_score
                    })

                    # Export options
                    csv_data = pd.DataFrame(results).to_csv(index=False)
                    st.download_button(
                        label="Download Results as CSV",
                        data=csv_data,
                        file_name=f"grading_results_{student_info['id']}.csv",
                        mime="text/csv"
                    )

        # Clean up temporary file
        if os.path.exists(file_path):
            os.remove(file_path)

    # Display all graded assignments
    if st.session_state.graded_assignments:
        st.subheader("All Graded Assignments")
        summary_data = []
        for assignment in st.session_state.graded_assignments:
            summary_data.append({
                'Student Name': assignment['student_info']['name'],
                'Student ID': assignment['student_info']['id'],
                'Rubric Score': f"{assignment['total_earned']:.1f}/{assignment['total_possible']:.1f}",
                'Rubric Percentage': f"{(assignment['total_earned']/assignment['total_possible'])*100:.1f}%",
                'Simple Score': f"{assignment['simple_score']}/80"
            })
        st.dataframe(pd.DataFrame(summary_data))

if __name__ == "__main__":
    main()
